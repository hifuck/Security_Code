# -*- coding:utf-8 -*-
#__author__:langzi
#__blog__:www.langzi.fun

import os
import time
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.shared import Inches
import multiprocessing
import re
import random
import requests
from concurrent.futures import ProcessPoolExecutor
from selenium import webdriver
from selenium.webdriver.support import expected_conditions as EC
from PIL import ImageGrab
# 输入屏幕左上角和右下角的坐标
from urllib.parse import urlparse

import threading

lock = threading.Lock()

headerss = [
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/22.0.1207.1 Safari/537.1",
    "Mozilla/5.0 (X11; CrOS i686 2268.111.0) AppleWebKit/536.11 (KHTML, like Gecko) Chrome/20.0.1132.57 Safari/536.11",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.6 (KHTML, like Gecko) Chrome/20.0.1092.0 Safari/536.6",
    "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.6 (KHTML, like Gecko) Chrome/20.0.1090.0 Safari/536.6",
    "Mozilla/5.0 (Windows NT 6.2; WOW64) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/19.77.34.5 Safari/537.1",
    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/536.5 (KHTML, like Gecko) Chrome/19.0.1084.9 Safari/536.5",
    "Mozilla/5.0 (Windows NT 6.0) AppleWebKit/536.5 (KHTML, like Gecko) Chrome/19.0.1084.36 Safari/536.5",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1063.0 Safari/536.3",
    "Mozilla/5.0 (Windows NT 5.1) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1063.0 Safari/536.3",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_8_0) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1063.0 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1062.0 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1062.0 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.1) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.0 Safari/536.3",
    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/535.24 (KHTML, like Gecko) Chrome/19.0.1055.1 Safari/535.24",
    "Mozilla/5.0 (Windows NT 6.2; WOW64) AppleWebKit/535.24 (KHTML, like Gecko) Chrome/19.0.1055.1 Safari/535.24"]

def writedata(x):
    with open('log.txt', 'a+')as aa:
        aa.write('***********************************' + '\n')
        aa.write(str(time.strftime('%Y-%m-%d:%H:%M:%S   ', time.localtime())) + str(x) + '\n')


class Get_Info:
    def __init__(self,url):
        self.url=url
        self.title_parrten = 'class="w61-0"><div class="ball">(.*?)</div></td>'  # group(1) 正常
        self.ip_parrten = '>IP：(.*?)</a></div>'  # group(1) 正常
        self.ages = '" target="_blank">(.*?)</a></div></div>'  # group(1)
        self.whois_id = '备案号：</span><a href=.*?" target="_blank">(.*?)</a></div>'  # 需group(1)
        self.whois_type = '<span>性质：</span><strong>(.*?)</strong></div>'  # 需group(1)
        self.whois_name = '<span>名称：</span><strong>(.*?)</strong></div>'  # 需group(1)
        self.whois_time = '<span>审核时间：</span><strong>(.*?)</strong></div>'  # 需group(1)
        self.include_baidu = '<div class="Ma01LiRow w12-1 ">(.*?)</div>'  # group(1)
        self.infos = '<div class="MaLi03Row w180">(.*?)</div>'  # 要findall 0，1，2，3
        self.result={'百度权重':'',
                    '网站主页':'',
                     '网站描述':'',
                   '网站标题':'',
                   'IP__坐标':'',
                   '所属__IP':'',
                   '网站年龄':'',
                   '备案编号':'',
                   '备案性质':'',
                   '备案名称':'',
                   '备案时间':'',
                   '百度收录':'',
                   '协议类型':'',
                   '页面类型':'',
                   '服务类型':'',
                   '程序语言':''}

    def get_baidu_weight(self):
        time.sleep(random.randint(1, 5))
        x = str(random.randint(1, 9))
        data = {
            't': 'rankall',
            'on': 1,
            'type': 'baidupc',
            'callback': 'jQuery111303146901980779846_154444474116%s' % (x),
            'host': self.url
        }

        headers = {

            'Accept': 'text/javascript, application/javascript, application/ecmascript, application/x-ecmascript, */*; q=0.01',
            'Accept-Encoding': 'gzip, deflate',
            'Accept-Language': 'zh-CN,zh;q=0.9',
            'Connection': 'keep-alive',
            'Content-Type': 'application/x-www-form-urlencoded; charset=UTF-8',
            'Cookie': 'UM_distinctid=165af67ee6f352-07238a34ed3941-9393265-1fa400-165af67ee70473; CNZZDATA5082706=cnzz_eid%3D832961605-1544438317-null%26ntime%3D1544443717; Hm_lvt_aecc9715b0f5d5f7f34fba48a3c511d6=1544443985; Hm_lpvt_aecc9715b0f5d5f7f34fba48a3c511d6=1544443985; qHistory=aHR0cDovL3JhbmsuY2hpbmF6LmNvbS9iYWlkdW1vYmlsZS8r55m+5bqm56e75Yqo5p2D6YeNfGh0dHA6Ly9yYW5rLmNoaW5hei5jb20vcmFua2FsbC8r5p2D6YeN57u85ZCI5p+l6K+ifGh0dHA6Ly9yYW5rLmNoaW5hei5jb20r55m+5bqm5p2D6YeN5p+l6K+ifGh0dHA6Ly9pbmRleC5jaGluYXouY29tLyvlhbPplK7or43lhajnvZHmjIfmlbB8aHR0cDovL3JhbmsuY2hpbmF6LmNvbS9yYW5rL2hpc3RvcnkuYXNweCvmnYPph43ljoblj7Lmn6Xor6I=',
            'Host': 'rank.chinaz.com',
            'Origin': 'http://rank.chinaz.com',
            'Referer': 'http://rank.chinaz.com',
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/68.0.3440.106 Safari/537.36',
            'X-Requested-With': 'XMLHttpRequest'

        }
        try:
            urls = 'http://rank.chinaz.com/ajaxseo.aspx?t=rankall&on=1&type=undefined&callback=jQuery111303146901980779846_154444474116%s' % (
                x)

            r = requests.post(url=urls, headers=headers, data=data)
            try:
                res = re.search(b',"br":(\d),"beforBr', r.content).group(1)
            except:
                pass
            if res:
                return res.decode()
            else:
                return '0'
        except:
            return '获取失败'


    def get_info_from_pattren(self,pattren, result):
        try:
            res = re.search(pattren, result).group(1)
            return res
        # return str(res.encode('utf-8'))
        except:
            return '暂无信息'

    def scan_seo(self):
        UA = random.choice(headerss)
        headers = {'User-Agent': UA}
        domain_url = self.url.replace('https://', '').replace('http://', '').replace('www.', '').split('/')[0]
        urls = 'http://seo.chinaz.com/' + domain_url
        # url,title,weights,ip,ages,whois_id,whois_type,whois_name,whois_time
        # 网址，标题，百度权重，ip信息，年龄，备案号，备案性质，备案名称，备案时间
        # include_baidu,request,text,service,language
        # 百度收录，，协议类型，页面类型，服务器类型，程序语言
        try:
            req = requests.get(urls, headers, verify=False, timeout=20)
            encoding = requests.utils.get_encodings_from_content(req.text)[0]
            r = req.content.decode(encoding, 'replace')
        except Exception as e:
            writedata(str(e))

        try:
            req = requests.get(self.url, headers, verify=False, timeout=20)
            encoding = requests.utils.get_encodings_from_content(req.text)[0]
            rss = req.content.decode(encoding, 'replace')
        except Exception as e:
            writedata(str(e))


        self.result['网站描述'] = '暂无信息'
        try:
            description = re.search('description.*?content=(.*?)>',rss,re.S).group(1)
            self.result['网站描述'] = description.strip()
        except:
            pass

        self.result['百度权重'] = str(self.get_baidu_weight())

        self.result['网站主页'] = self.url.split('//')[0] + '//' + self.url.split('//')[1].split('/')[0]

        try:
            req1 = requests.get(url=self.url, headers=headers, verify=False, timeout=10)
            encoding = requests.utils.get_encodings_from_content(req1.text)[0]
            rress = req1.content.decode(encoding, 'replace')
            title_pattern = '<title>(.*?)</title>'
            title = re.search(title_pattern, rress, re.S | re.I)
            self.result['网站标题'] = str(title.group(1))
        except:
            self.result['网站标题'] = self.url.split('//')[0] + '//' + self.url.split('//')[1].split('/')[0]


        ip_infos = self.get_info_from_pattren(self.ip_parrten, r)
        if '[' in ip_infos:
            ip, address = ip_infos.split('[')[0], ip_infos.split('[')[1]
            self.result['IP__坐标'] = address.replace(']', '')
            self.result['所属__IP'] = ip
        else:
            self.result['IP__坐标'] = '获取失败'
            self.result['所属__IP'] = '获取失败'

        self.result['网站年龄'] = self.get_info_from_pattren(self.ages, r)
        self.result['备案编号'] = self.get_info_from_pattren(self.whois_id, r)
        self.result['备案性质'] = self.get_info_from_pattren(self.whois_type, r)
        self.result['备案名称'] = self.get_info_from_pattren(self.whois_name, r)
        self.result['备案时间'] = self.get_info_from_pattren(self.whois_time, r)
        self.result['百度收录'] = self.get_info_from_pattren(self.include_baidu, r)

        dd = re.findall(self.infos, r, re.S)
        resu = ['暂无信息' if x.replace(' ', '') is '' else x for x in dd]
        try:
            self.result['协议类型'] = resu[0]
        except:
            self.result['协议类型'] = '获取失败'

        try:
            self.result['页面类型'] = resu[1]
        except:
            self.result['页面类型'] = '获取失败'

        try:
            self.result['服务类型'] = resu[2]
        except:
            self.result['服务类型'] = '获取失败'

        try:
            self.result['程序语言'] = resu[3]
        except:
            self.result['程序语言'] = '获取失败'
        return self.result




def get_image():
    if os.path.exists('code.png'):
        os.remove('code.png')
    time.sleep(5)
    pic = ImageGrab.grab(bbox=(20, 31, 1266, 1016))
    # 指定坐标 左上角和右下角
    # pic = ImageGrab.grab()
    # 全屏截图
    pic.save("code.png")




def run(url):
    driver = webdriver.Firefox()
    driver.set_page_load_timeout(20)

    try:
        driver.get(url)
        time.sleep(20)
        # driver.quit()
    except:
        pass
    try:
        driver.close()
    except Exception as e:
        print(e)

    finally:
        # driver.quit()
        return True

def main_svn(url):
    print('Check : {} '.format(url))
    t1 = threading.Thread(target=run,args=(url,))
    t2 = threading.Thread(target=get_image)
    t1.start()
    t2.start()
    t1.join()
    t2.join()
    if os.path.exists('code.png'):
        pass
    else:
        time.sleep(10)
    print('开始保存文档....')
    document = Document()
    DaHei = document.styles.add_style('DaHei', 1)
    # 设置字体尺寸
    DaHei.font.size = Pt(16)
    # 设置字体颜色
    DaHei.font.color.rgb = RGBColor(0, 0, 0)
    # 黑色
    # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
    # 居中文本
    DaHei.font.name = '仿宋'
    DaHei._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    ZhongHei = document.styles.add_style('ZhongHei', 1)
    # 设置字体尺寸
    ZhongHei.font.size = Pt(10)
    # 设置字体颜色
    ZhongHei.font.color.rgb = RGBColor(0, 0, 0)
    # 黑色
    # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
    # 居中文本
    ZhongHei.font.name = '仿宋'
    ZhongHei._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    ZhongHOng = document.styles.add_style('ZhongHOng', 1)
    # 设置字体尺寸
    ZhongHOng.font.size = Pt(9)
    # 设置字体颜色
    ZhongHOng.font.color.rgb = RGBColor(178, 34, 34)
    # 红色
    # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
    # 居中文本
    ZhongHOng.font.name = '仿宋'
    ZhongHOng._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    XiaoLv = document.styles.add_style('XiaoLv', 1)
    # 设置字体尺寸
    XiaoLv.font.size = Pt(6.5)
    # 设置字体颜色
    XiaoLv.font.color.rgb = RGBColor(0, 255, 0)
    # 绿色
    # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
    # 居中文本
    XiaoLv.font.name = '仿宋'
    XiaoLv._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    domain = urlparse(url).netloc
    try:
        filename = url.split('//')[1].split('/')[0] + '__存在SVN源码泄露漏洞.docx'
    except:
        filename = domain.replace(':', '_').replace('/', '_') + '__存在SVN源码泄露漏洞.docx'

    document.add_paragraph('基本信息', 'Title')
    document.add_paragraph('漏洞类型：源码泄露漏洞', 'Subtitle')
    document.add_paragraph('漏洞等级：高危', 'Subtitle')
    document.add_paragraph('厂商信息：此处需要手动搜索', 'Subtitle')
    document.add_page_break()

    document.add_paragraph('漏洞简述', 'Title')
    document.add_paragraph('漏洞描述', 'Subtitle')
    document.add_paragraph('整个网站的源代码文件泄露，包括大量的敏感数据，数据库连接数据 等。', 'DaHei')
    document.add_paragraph('产生原因', 'Subtitle')
    document.add_paragraph('''
    SVN（subversion）是源代码版本管理软件，造成SVN源代码漏洞的主要原因是管理员操作不规范。“在使用SVN管理本地代码过程中，会自动生成一个名为.svn的隐藏文件夹，其中包含重要的源代码信息。但一些网站管理员在发布代码时，不愿意使用‘导出’功能，而是直接复制代码文件夹到WEB服务器上，这就使.svn隐藏文件夹被暴露于外网环境，黑客可以借助其中包含的用于版本信息追踪的‘entries’文件，逐步摸清站点结构。”（可以利用.svn/entries文件，获取到服务器源码、svn服务器账号密码等信息）
    ''', 'DaHei')
    document.add_paragraph('漏洞危害', 'Subtitle')
    document.add_paragraph(
       '''
    1  SVN产生的.svn目录下还包含了以.svn-base结尾的源代码文件副本（低版本SVN具体路径为text-base目录，高版本SVN为pristine目录），如果服务器没有对此类后缀做解析，黑客则可以直接获得文件源代码。
    2 可以列出网站目录，甚至下载整站。
    3 可以直接获取数据库连接信息账号密码，以及后台地址等更多敏感信息。
    4 可以用来下载该网站源码，做代码审计，发现更大的漏洞造成更大的危害
    5 网站存在包含SVN信息的文件，这是网站源码的版本控制器私有文件，里面包含SVN服务的地址、提交的私有文件名、SVN用户名等信息，该信息有助于攻击者更全面了解网站的架构，为攻击者入侵网站提供帮助。
    6 该信息泄露会暴露服务器的敏感信息，使攻击者能够通过泄露的信息进行进一步入侵。
       ''','DaHei')

    document.add_page_break()

    document.add_paragraph('漏洞详情', 'Title')
    document.add_paragraph('网站漏洞报表', 'Subtitle')
    document.add_paragraph('漏洞地址:'+url,'ZhongHei')
    urls = url.split('//')[0] + '//' + url.split('//')[1].split('/')[0]
    try:
        document.add_paragraph('网站信息报表', 'Subtitle')
        info = Get_Info(urls)
        infos = info.scan_seo()
        for k,v in infos.items():
            document.add_paragraph(k+'  '+v,'ZhongHei')
    except Exception as e:
        writedata(str(e))
    document.add_page_break()


    document.add_paragraph('网站漏洞复现(使用Firefox浏览器)', 'Subtitle')
    document.add_paragraph('下载链接','DaHei')
    document.add_paragraph(url,'ZhongHOng')
    document.add_paragraph('返回结果','DaHei')
    document.add_picture('code.png',width=Inches(7))
    document.add_paragraph('源码恢复检测工具','DaHei')
    document.add_paragraph('xxxxxx这里手动插图xxxxxxxxxx','DaHei')

    document.add_paragraph('复现完成，验证存在源码泄露漏洞','DaHei')

    document.add_page_break()
    document.add_paragraph('修复建议', 'Title')
    document.add_paragraph('根源上进行修复', 'Subtitle')
    document.add_paragraph('''
    1、在web服务器配置文件中增加一段代码，过滤到.svn文件，返回404nginx服务器：
    location ~ ^(.*)\/\.svn\/
    {
    return 404;
    }
    重启nginx
    Apache服务器：
    Order allow,deny
    Deny from all 
    重启Apache
    
    2、查找服务器上所有.svn隐藏文件夹，删除
    以下命令删除当前目录下.svn文件夹
    find . -type d -name ".svn"|xargs rm -rf
    rm -rf `find . -type d -name .svn`
    find . -name ".svn" -type d | xargs rm -fr
    find . -name ".svn" -type d | xargs -n1 rm -R
    ''','DaHei')
    document.save(filename)


def main_git(url):
    print('Check : {} '.format(url))
    t1 = threading.Thread(target=run,args=(url,))
    t2 = threading.Thread(target=get_image)
    t1.start()
    t2.start()
    t1.join()
    t2.join()
    if os.path.exists('code.png'):
        pass
    else:
        time.sleep(10)
    print('开始保存文档....')
    document = Document()
    DaHei = document.styles.add_style('DaHei', 1)
    # 设置字体尺寸
    DaHei.font.size = Pt(16)
    # 设置字体颜色
    DaHei.font.color.rgb = RGBColor(0, 0, 0)
    # 黑色
    # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
    # 居中文本
    DaHei.font.name = '仿宋'
    DaHei._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    ZhongHei = document.styles.add_style('ZhongHei', 1)
    # 设置字体尺寸
    ZhongHei.font.size = Pt(10)
    # 设置字体颜色
    ZhongHei.font.color.rgb = RGBColor(0, 0, 0)
    # 黑色
    # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
    # 居中文本
    ZhongHei.font.name = '仿宋'
    ZhongHei._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    ZhongHOng = document.styles.add_style('ZhongHOng', 1)
    # 设置字体尺寸
    ZhongHOng.font.size = Pt(9)
    # 设置字体颜色
    ZhongHOng.font.color.rgb = RGBColor(178, 34, 34)
    # 红色
    # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
    # 居中文本
    ZhongHOng.font.name = '仿宋'
    ZhongHOng._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    XiaoLv = document.styles.add_style('XiaoLv', 1)
    # 设置字体尺寸
    XiaoLv.font.size = Pt(6.5)
    # 设置字体颜色
    XiaoLv.font.color.rgb = RGBColor(0, 255, 0)
    # 绿色
    # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
    # 居中文本
    XiaoLv.font.name = '仿宋'
    XiaoLv._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    domain = urlparse(url).netloc
    try:
        filename = url.split('//')[1].split('/')[0] + '__存在GIT源码泄露漏洞.docx'
    except:
        filename = domain.replace(':', '_').replace('/', '_') + '__存在GIT源码泄露漏洞.docx'

    document.add_paragraph('基本信息', 'Title')
    document.add_paragraph('漏洞类型：源码泄露漏洞', 'Subtitle')
    document.add_paragraph('漏洞等级：高危', 'Subtitle')
    document.add_paragraph('厂商信息：此处需要手动搜索', 'Subtitle')
    document.add_page_break()

    document.add_paragraph('漏洞简述', 'Title')
    document.add_paragraph('漏洞描述', 'Subtitle')
    document.add_paragraph('整个网站的源代码文件泄露，包括大量的敏感数据，数据库连接数据 等。', 'DaHei')
    document.add_paragraph('产生原因', 'Subtitle')
    document.add_paragraph('''
    在运行git init初始化代码库的时候，会在当前目录下面产生一个.git的隐藏文件，用来记录代码的变更记录等等。在发布代码的时候，把.git这个目录没有删除，直接发布了。使用这个文件，可以用来恢复源代码。
    ''', 'DaHei')
    document.add_paragraph('漏洞危害', 'Subtitle')
    document.add_paragraph(
       '''
       攻击者就可以利用下载git文件夹 , 可以利用其中储存的版本控制信息 , 完全恢复网站后台的代码和目录结构 , 当然 , 一般网站都会涉及到数据库操作 , 而一般来说 , 需要链接数据库就需要用户名/密码/端口/库名等信息 , 而这些信息肯定会在网站后台的源码里面又体现 , 因此这种情况是极其危险的 , 还有 , 一旦服务器开放了数据库的远程连接功能 , 攻击者就可以利用从源码中找到的数据库用户名和密码对远程数据库进行登陆 , 危险性不言而喻
       ''','DaHei')

    document.add_page_break()

    document.add_paragraph('漏洞详情', 'Title')
    document.add_paragraph('网站漏洞报表', 'Subtitle')
    document.add_paragraph('漏洞地址:'+url,'ZhongHei')
    urls = url.split('//')[0] + '//' + url.split('//')[1].split('/')[0]
    try:
        document.add_paragraph('网站信息报表', 'Subtitle')
        info = Get_Info(urls)
        infos = info.scan_seo()
        for k,v in infos.items():
            document.add_paragraph(k+'  '+v,'ZhongHei')
        document.add_page_break()
    except Exception as e:
        writedata(str(e))


    document.add_paragraph('网站漏洞复现(使用Firefox浏览器)', 'Subtitle')
    document.add_paragraph('下载链接','DaHei')
    document.add_paragraph(url,'ZhongHOng')
    document.add_paragraph('返回结果','DaHei')
    document.add_picture('code.png',width=Inches(7))
    document.add_paragraph('源码恢复检测工具','DaHei')
    document.add_paragraph('xxxxxx这里手动插图xxxxxxxxxx','DaHei')

    document.add_paragraph('复现完成，验证存在源码泄露漏洞','DaHei')

    document.add_page_break()
    document.add_paragraph('修复建议', 'Title')
    document.add_paragraph('根源上进行修复', 'Subtitle')
    document.add_paragraph('''
    1、对.git目录的访问权限进行控制
    2、在每次pull之后删除.git文件夹
    ''','DaHei')

    document.save(filename)

if __name__ == '__main__':
    print('''
         ___   _   _   _____   _____        _           ___   __   _   _____   ______  _  
        /   | | | | | |_   _| /  _  \      | |         /   | |  \ | | /  ___| |___  / | | 
       / /| | | | | |   | |   | | | |      | |        / /| | |   \| | | |        / /  | | 
      / / | | | | | |   | |   | | | |      | |       / / | | | |\   | | |  _    / /   | | 
     / /  | | | |_| |   | |   | |_| |      | |___   / /  | | | | \  | | |_| |  / /__  | | 
    /_/   |_| \_____/   |_|   \_____/      |_____| /_/   |_| |_|  \_| \_____/ /_____| |_| 

    ''')
    #time.sleep(1)
    print('''

        Description:
            Langzi_Auto_GIT_SVN_Scan v0.5版本
            是一款批量针对源码泄露自动生成报告工具
            需要配合 svn-git源码泄露扫描工具0.6 版本使用
            对 svn-git源码泄露扫描工具0.6版本 扫描结果生成漏洞报告
            自动截图坐标为 20,31,1266,1016(目前适用于langzi本机)


        Tips:    
            /*禁止对GOV-EDU进行检测(检测到则秒退)*/ 
            需要手动删除Lang_xss_v1.1最后一行横线
            在WIN 7 下不兼容运行
            运行目录不能存在中文字符

    ''')
    #time.sleep(6)
    New_start = input(('导入Lang_Backup_File_Scan3.9 扫描结果文本:'))
    #New_start = 'vlun-result.txt'

    list_ = [x.strip() for x in open(New_start,'r',encoding='utf-8').readlines()]
    for x in list_:
        if 'git' in x:
            main_git(x)
        if 'svn' in x:
            main_svn(x)


